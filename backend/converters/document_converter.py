import os
import logging
import html as html_mod
import re
from typing import Dict, List

from .base import BaseConverter

logger = logging.getLogger("hidconverter")

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
except ImportError:
    DocxDocument = None
    qn = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

try:
    import markdown
except ImportError:
    markdown = None

_UNICODE_FONT_PATHS = [
    ("C:/Windows/Fonts/arial.ttf", "Arial"),
    ("C:/Windows/Fonts/DejaVuSans.ttf", "DejaVuSans"),
    ("C:/Windows/Fonts/NotoSans-Regular.ttf", "NotoSans"),
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVuSans"),
    ("/usr/share/fonts/TTF/DejaVuSans.ttf", "DejaVuSans"),
    ("/System/Library/Fonts/Supplemental/Arial.ttf", "Arial"),
    ("/Library/Fonts/Arial.ttf", "Arial"),
]

MAX_WORD_LENGTH = 70


def _get_unicode_font():
    for path, name in _UNICODE_FONT_PATHS:
        if os.path.exists(path):
            return path, name
    return None, None


class DocumentConverter(BaseConverter):
    SUPPORTED = {
        "pdf":  ["docx", "txt"],
        "docx": ["pdf"],
        "txt":  ["pdf"],
        "md":   ["pdf"],
        "xlsx": ["csv", "pdf"],
    }

    def supported_conversions(self) -> Dict[str, List[str]]:
        return self.SUPPORTED

    def convert(self, file_path: str, target_format: str, output_dir: str) -> str:
        ext = os.path.splitext(file_path)[1].lstrip(".").lower()
        if ext not in self.SUPPORTED or target_format not in self.SUPPORTED[ext]:
            raise ValueError(f"Conversion from {ext} to {target_format} is not supported")

        output_path = self.get_output_path(file_path, target_format, output_dir)
        func_name = f"_convert_{ext}_to_{target_format}"
        handler = getattr(self, func_name, None)
        if not handler:
            raise ValueError(f"No handler for {ext} -> {target_format}")

        return handler(file_path, output_path)

    # --- helpers ---

    def _new_pdf(self):
        if FPDF is None:
            raise RuntimeError("Missing dependency: fpdf2")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        font_path, font_name = _get_unicode_font()
        if font_path:
            pdf.add_font(font_name, "", font_path, uni=True)
            pdf.set_font(font_name, size=11)
        else:
            pdf.set_font("Helvetica", size=11)
        return pdf

    def _safe_text(self, text):
        words = text.split(" ")
        safe_words = []
        for word in words:
            while len(word) > MAX_WORD_LENGTH:
                safe_words.append(word[:MAX_WORD_LENGTH])
                word = word[MAX_WORD_LENGTH:]
            safe_words.append(word)
        return " ".join(safe_words)

    def _write_paragraph(self, pdf, text):
        text = text or ""
        safe = self._safe_text(text)
        available = pdf.w - pdf.l_margin - pdf.r_margin
        try:
            pdf.multi_cell(w=available, h=6, text=safe)
        except Exception as e:
            logger.warning(f"Skipping paragraph: {e}")

    def _write_line(self, pdf, text=""):
        if text.strip():
            self._write_paragraph(pdf, text)
        else:
            pdf.ln(6)

    def _write_table_row(self, pdf, cells):
        available = pdf.w - pdf.l_margin - pdf.r_margin
        cols = max(len(cells), 1)
        col_w = available / cols
        for cell in cells:
            text = self._safe_text(str(cell or ""))
            try:
                pdf.multi_cell(w=col_w, h=6, text=text)
            except Exception as e:
                logger.warning(f"Skipping table cell: {e}")

    def _render_docx_body(self, doc, pdf):
        if qn is None:
            raise RuntimeError("Missing dependency: python-docx")

        body = doc.element.body
        for child in body:
            tag = child.tag
            if tag.endswith("tbl"):
                try:
                    from docx.table import Table
                    table = Table(child, doc)
                    for row in table.rows:
                        cells = [c.text for c in row.cells]
                        self._write_table_row(pdf, cells)
                        pdf.ln(2)
                except Exception as e:
                    logger.warning(f"Skipping table: {e}")
            elif tag.endswith("p"):
                try:
                    from docx.text.paragraph import Paragraph
                    p = Paragraph(child, doc)
                    self._write_line(pdf, p.text)
                except Exception as e:
                    logger.warning(f"Skipping paragraph: {e}")

    # --- PDF -> DOCX ---

    def _convert_pdf_to_docx(self, file_path: str, output_path: str) -> str:
        if pdfplumber is None or DocxDocument is None:
            raise RuntimeError("Missing dependencies: pdfplumber or python-docx")

        doc = DocxDocument()
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                doc.add_paragraph(text)
        doc.save(output_path)
        return output_path

    # --- PDF -> TXT ---

    def _convert_pdf_to_txt(self, file_path: str, output_path: str) -> str:
        if pdfplumber is None:
            raise RuntimeError("Missing dependency: pdfplumber")

        with pdfplumber.open(file_path) as pdf:
            text = "\n\n".join(page.extract_text() or "" for page in pdf.pages)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        return output_path

    # --- DOCX -> PDF ---

    def _convert_docx_to_pdf(self, file_path: str, output_path: str) -> str:
        if DocxDocument is None:
            raise RuntimeError("Missing dependency: python-docx")

        doc = DocxDocument(file_path)
        pdf = self._new_pdf()
        self._render_docx_body(doc, pdf)
        pdf.output(output_path)
        return output_path

    # --- TXT -> PDF ---

    def _convert_txt_to_pdf(self, file_path: str, output_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        pdf = self._new_pdf()
        for line in text.split("\n"):
            self._write_line(pdf, line)
        pdf.output(output_path)
        return output_path

    # --- MD -> PDF ---

    def _convert_md_to_pdf(self, file_path: str, output_path: str) -> str:
        if markdown is None:
            raise RuntimeError("Missing dependency: markdown")

        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        html = markdown.markdown(text)
        plain = html_mod.unescape(re.sub(r"<[^>]+>", "", html))

        pdf = self._new_pdf()
        for line in plain.split("\n"):
            self._write_line(pdf, line.strip())
        pdf.output(output_path)
        return output_path

    # --- XLSX -> CSV ---

    def _convert_xlsx_to_csv(self, file_path: str, output_path: str) -> str:
        if openpyxl is None:
            raise RuntimeError("Missing dependency: openpyxl")

        import csv
        wb = openpyxl.load_workbook(file_path, read_only=True)
        ws = wb.active
        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(row)
        return output_path

    # --- XLSX -> PDF ---

    def _convert_xlsx_to_pdf(self, file_path: str, output_path: str) -> str:
        if openpyxl is None:
            raise RuntimeError("Missing dependency: openpyxl")

        wb = openpyxl.load_workbook(file_path, read_only=True)
        ws = wb.active
        pdf = self._new_pdf()
        available = pdf.w - pdf.l_margin - pdf.r_margin
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell or "") for cell in row]
            cols = max(len(cells), 1)
            col_w = available / min(cols, 6)
            num_groups = max(1, cols // 6 + (1 if cols % 6 else 0))
            for g in range(num_groups):
                group = cells[g * 6:(g + 1) * 6]
                for cell in group:
                    text = self._safe_text(cell)
                    try:
                        pdf.multi_cell(w=col_w, h=5, text=text)
                    except Exception as e:
                        logger.warning(f"Skipping xlsx cell: {e}")
                pdf.ln(5)
        pdf.output(output_path)
        return output_path
